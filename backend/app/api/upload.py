from fastapi import APIRouter, File, UploadFile, Depends, HTTPException
from sqlalchemy.orm import Session
from app.core.database import get_db, SessionLocal
from app.models.contract import Contract, ContractIssue
from app.services.document_parser import parse_document
from app.services.chunking import processing_queue, chunk_text
from app.services.llm_pipeline import extract_financials, analyze_chunk_with_llm
from app.api.websockets import manager
import asyncio
import datetime

router = APIRouter()

async def process_document_background(contract_id: int, parsed_text: str):
    current_db = SessionLocal()
    try:
        # 1. Extract Financials quickly
        print(f"Starting background processing for contract {contract_id}")
        financials = await extract_financials(parsed_text)
        print(f"Extracted {len(financials)} financials")
        
        # Sort out the key business terms to highlight
        business_terms = {
            "contract_value": "Not Found",
            "payment_terms": "Not Found",
            "retainage": "Not Found"
        }
        other_financials = []
        
        financials_list = financials if isinstance(financials, list) else [] # type bypass
        # For actual iterable financials
        if isinstance(financials, list):
            for fin in financials:
                if not isinstance(fin, dict): continue
                label = str(fin.get("item", "")).lower()
                if "contract sum" in label or "contract value" in label or "total amount" in label:
                    business_terms["contract_value"] = fin.get("amount", "Not Found")
                elif "payment" in label or "net" in label or "terms" in label:
                    business_terms["payment_terms"] = fin.get("amount", "Not Found")
                elif "retention" in label or "retainage" in label:
                    business_terms["retainage"] = fin.get("amount", "Not Found")
                else:
                    other_financials.append(fin)
                    
        business_terms["other"] = other_financials
                
        contract = current_db.query(Contract).filter(Contract.id == contract_id).first()
        if contract:
            contract.financials = business_terms  # Save the structured dictionary
            current_db.commit()
        
        # 2. Chunk text and add to rate-limited queue
        chunks = chunk_text(parsed_text)
        total_chunks = len(chunks)
        print(f"Created {total_chunks} text chunks for processing")
        
        # Track progress
        progress_state = {"completed": 0, "total": total_chunks}
        
        # Broadcast initial state
        await manager.broadcast_progress(contract_id, 0, total_chunks)
        
        async def process_single_chunk(c_text, cid):
            print(f"Processing chunk for contract {cid}...")
            result = await analyze_chunk_with_llm(c_text)
            
            issues = result.get("issues", [])
            summary = result.get("summary", {})
            
            if summary:
                contract_db = SessionLocal()
                try:
                    import json
                    contract_obj = contract_db.query(Contract).filter(Contract.id == cid).first()
                    if contract_obj:
                        current_summary = json.loads(contract_obj.summary_text) if contract_obj.summary_text else {"sections": [], "exhibits": []}
                        
                        new_sections = summary.get("sections", [])
                        new_exhibits = summary.get("exhibits", [])
                        
                        if isinstance(new_sections, list):
                            current_summary["sections"].extend(new_sections)
                        if isinstance(new_exhibits, list):
                            current_summary["exhibits"].extend(new_exhibits)
                            
                        contract_obj.summary_text = json.dumps(current_summary)
                        contract_db.commit()
                except Exception as e:
                    print(f"Failed to update summary: {e}")
                finally:
                    contract_db.close()
            
            if issues:
                print(f"Found {len(issues)} issues in chunk, broadcasting...")
                
                # Save to database in a localized background session
                chunk_db = SessionLocal()
                try:
                    for issue in issues:
                        options = issue.get("alternative_language_options", [])
                        option_a = options[0].get("language", "") if len(options) > 0 else ""
                        option_b = options[1].get("language", "") if len(options) > 1 else ""
                        is_deletion = issue.get("proposed_deletion", False)

                        # Map the new prompt JSON to our standard DB/Frontend schema
                        mapped_issue = {
                            "contract_id": cid,
                            "category": issue.get("section_name", issue.get("section", "General")),
                            "original_text": issue.get("original_text", ""),
                            "location": f"Page {issue.get('page', 'Unknown')}",
                            "risk_level": "High" if is_deletion else "Medium",
                            "explanation": issue.get("problem_reasoning", ""),
                            "option_a_text": "[DELETE ENTIRELY]" if is_deletion else option_a,
                            "option_b_text": option_b
                        }

                        db_issue = ContractIssue(**mapped_issue)
                        chunk_db.add(db_issue)
                        
                        # Broadcast the mapped version so the Frontend React component doesn't break
                        await manager.broadcast_issue(cid, mapped_issue)
                    chunk_db.commit()
                except Exception as e:
                    print(f"Failed to save issues to DB: {e}")
                    chunk_db.rollback()
                    chunk_db.close()
                
            progress_state["completed"] += 1
            await manager.broadcast_progress(cid, progress_state["completed"], progress_state["total"])
                
        for chunk in chunks:
            await processing_queue.add_task(process_single_chunk, chunk, contract_id)
    except Exception as e:
        print(f"CRITICAL FASTAPI BACKGROUND ERROR: {e}")
    finally:
        current_db.close()

@router.post("/upload/")
async def upload_contract(file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename.endswith(('.pdf', '.docx', '.txt')):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, DOCX, and TXT are supported.")
    
    file_bytes = await file.read()
    
    # 1. Parse text
    try:
        parsed_text = parse_document(file.filename, file_bytes)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")
        
    # 2. Save Contract record to DB
    new_contract = Contract(
        filename=file.filename,
        status="processing",
        uploaded_at=datetime.datetime.utcnow()
    )
    db.add(new_contract)
    db.commit()
    db.refresh(new_contract)
    
    # Trigger background LLM pipeline
    asyncio.create_task(process_document_background(new_contract.id, parsed_text))
    
    return {
        "message": "File uploaded and processing started in background.", 
        "contract_id": new_contract.id,
        "filename": new_contract.filename,
        "parsed_length": len(parsed_text),
        "parsed_text": parsed_text
    }

from fastapi.responses import StreamingResponse
import io
from docx import Document
from docx.shared import RGBColor

@router.get("/history")
async def get_contract_history(db: Session = Depends(get_db)):
    contracts = db.query(Contract).order_by(Contract.uploaded_at.desc()).all()
    return [{"id": c.id, "filename": c.filename, "uploaded_at": c.uploaded_at, "status": c.status} for c in contracts]

@router.delete("/{contract_id}")
async def delete_contract(contract_id: int, db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    # Delete associated issues to keep DB clean
    db.query(ContractIssue).filter(ContractIssue.contract_id == contract_id).delete()
    db.delete(contract)
    db.commit()
    return {"status": "success", "message": "Contract permanently removed"}

@router.get("/{contract_id}/export")
async def export_contract(contract_id: int, db: Session = Depends(get_db)):
    # 1. Fetch Contract from DB
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
        
    try:
        # Create a new Word Document for MVP
        doc = Document()
        doc.add_heading(f"Scaffold AI Review: {contract.filename}", 0)
        
        doc.add_heading("Financial Summary", level=1)
        if contract.financials and isinstance(contract.financials, dict):
            if contract.financials.get('contract_value') and contract.financials.get('contract_value') != 'Not Found':
                doc.add_paragraph(f"Total Contract Value: {contract.financials.get('contract_value')}")
            if contract.financials.get('payment_terms') and contract.financials.get('payment_terms') != 'Not Found':
                doc.add_paragraph(f"Payment Terms: {contract.financials.get('payment_terms')}")
            if contract.financials.get('retainage') and contract.financials.get('retainage') != 'Not Found':
                doc.add_paragraph(f"Retainage: {contract.financials.get('retainage')}")
            
            other = contract.financials.get('other', [])
            for fin in other:
                doc.add_paragraph(f"{fin.get('item', 'Item')}: {fin.get('amount', '$0')}")
        else:
            doc.add_paragraph("No financials identified.")
            
        doc.add_heading("Identified Risks & Suggested Redlines", level=1)
        doc.add_paragraph("The following risks were identified based on your scaffolding directives:")
        
        issues = db.query(ContractIssue).filter(ContractIssue.contract_id == contract_id).all()
        
        if issues:
            for issue in issues:
                doc.add_heading(f"Risk: {issue.category}", level=2)
                p = doc.add_paragraph()
                p.add_run(f"Original Text ({issue.location}): ").bold = True
                p.add_run(issue.original_text)
                
                p2 = doc.add_paragraph()
                p2.add_run("Explanation: ").bold = True
                p2.add_run(issue.explanation)
                
                p3 = doc.add_paragraph()
                p3.add_run("Option A (Most Protective): ").bold = True
                p3.add_run(issue.option_a_text).font.color.rgb = RGBColor(255, 0, 0)
                
                if issue.option_b_text and issue.option_b_text != "":
                    p4 = doc.add_paragraph()
                    p4.add_run("Option B (Compromise): ").bold = True
                    p4.add_run(issue.option_b_text).font.color.rgb = RGBColor(0, 0, 255)
        else:
            doc.add_paragraph("\nNo issues were identified in this contract.")
        
        # Save document to an in-memory byte stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Return as a downloadable file response
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=AI_Reviewed_{contract.filename}.docx"
            }
        )
    except Exception as e:
        print(f"Export Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate Word document.")

@router.get("/{contract_id}/export_summary")
async def export_contract_summary(contract_id: int, db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
        
    try:
        import json
        doc = Document()
        doc.add_heading(f"Contract Cheat Sheet: {contract.filename}", 0)
        
        doc.add_heading("Financial & Business Terms", level=1)
        if contract.financials:
            if contract.financials.get('contract_value') and contract.financials.get('contract_value') != 'Not Found':
                p = doc.add_paragraph()
                p.add_run("Total Contract Value: ").bold = True
                p.add_run(str(contract.financials.get('contract_value')))
            if contract.financials.get('payment_terms') and contract.financials.get('payment_terms') != 'Not Found':
                p = doc.add_paragraph()
                p.add_run("Payment Terms: ").bold = True
                p.add_run(str(contract.financials.get('payment_terms')))
            if contract.financials.get('retainage') and contract.financials.get('retainage') != 'Not Found':
                p = doc.add_paragraph()
                p.add_run("Retainage: ").bold = True
                p.add_run(str(contract.financials.get('retainage')))
            
            other = contract.financials.get('other', [])
            if other:
                doc.add_heading("Other Financials", level=2)
                for f in other:
                    doc.add_paragraph(f"{f.get('item', 'Item')}: {f.get('amount', '$0')}")
        else:
            doc.add_paragraph("No financials identified.")
            
        doc.add_heading("Sections Summary", level=1)
        summary_data = {"sections": [], "exhibits": []}
        if contract.summary_text:
            try:
                summary_data = json.loads(contract.summary_text)
            except:
                pass
                
        sections = summary_data.get("sections", [])
        if sections:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Section No.'
            hdr_cells[1].text = 'Title'
            hdr_cells[2].text = 'Page(s)'
            
            for sec in sections:
                row_cells = table.add_row().cells
                row_cells[0].text = str(sec.get("number", ""))
                row_cells[1].text = str(sec.get("title", ""))
                
                page_text = str(sec.get("start_page", ""))
                if sec.get("end_page") and sec.get("end_page") != sec.get("start_page"):
                    page_text += f" - {sec.get('end_page')}"
                elif sec.get("page"): 
                    page_text = str(sec.get("page"))
                row_cells[2].text = page_text
        else:
            doc.add_paragraph("No sections were extracted.")
            
        doc.add_heading("Exhibits & Attachments", level=1)
        exhibits = summary_data.get("exhibits", [])
        if exhibits:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Exhibit Label'
            hdr_cells[1].text = 'Title/Description'
            hdr_cells[2].text = 'Page'
            
            for ex in exhibits:
                row_cells = table.add_row().cells
                row_cells[0].text = str(ex.get("label", ""))
                row_cells[1].text = str(ex.get("title", ""))
                row_cells[2].text = str(ex.get("page", ""))
        else:
            doc.add_paragraph("No exhibits were extracted.")
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=CheatSheet_{contract.filename}.docx"
            }
        )
    except Exception as e:
        print(f"Export Summary Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate Cheat Sheet document.")
